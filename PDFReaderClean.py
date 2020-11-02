# -*- coding: utf-8 -*-
"""
Created on Mon Nov  2 10:45:30 2020

@author: 014174

PDF Reader - clean
"""

# =============================================================================
# Package import
# =============================================================================

import os 
import pandas as pd
import numpy as np
import PyPDF2 as p2
import docx
import collections
from PIL import Image 
import pytesseract 
import sys 
from pdf2image import convert_from_path 

# =============================================================================
# Path specification 
# =============================================================================
main_path = ''
pytesseract.pytesseract.tesseract_cmd = ''
poppler_path = main_path + "/poppler-0.68.0/bin"

# =============================================================================
# master function
# =============================================================================

def master(original_file):
    import docx
    import PyPDF2 as p2
    

    if original_file.endswith('.docx'):
        doc = docx.Document(original_file)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        tables = doc.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text.append(paragraph.text)
                        
        full_text = ' '.join(full_text)
            
            
        return(full_text)
            
    else:
        
        full_text = str()
        PDF_file = open(original_file, "rb")
        PDF_read = p2.PdfFileReader(PDF_file)
        file_length = PDF_read.getNumPages()
        ## Note: This is currently quite lazily implemented. If the file is a pdf it is converted into a poppler object regardless of whether that is required or not. This will slow execution down
        pages = convert_from_path(original_file, dpi = 500, poppler_path = poppler_path) 
        i = 0

        while i < file_length:
            page_number = PDF_read.getPage(i)
            page_text = page_number.extractText()
            
            if not page_text:
                
                pages[i].save(main_path + "/interim.png", 'PNG')
                page_text = str(pytesseract.image_to_string(Image.open(main_path + "/interim.png"))) 
                
                if len(page_text) > 0:
                    full_text = full_text + " " + page_text
                i += 1
                return(full_text)
                
                
            else:
                full_text = full_text + " " + page_text
                i += 1
        
                return(full_text)
            
# =============================================================================
# Tokenize, lemmatize and vectorize
# =============================================================================

def token_lem_vector(full_text):
    
    import pandas as pd
    
      
    if not full_text:
    
        return(pd.DataFrame())
        
    else:
        
        import nltk
        import re
        text = nltk.word_tokenize(full_text)
        from nltk import bigrams
    
    
        from nltk.corpus import stopwords
        nltk.download('stopwords')
        stop_words = set(stopwords.words("English"))
        text = [x for x in text if x != ' ']
        text = [x for x in text if x != ' ']
        text = [x for x in text if x != '']
        text = [x for x in text if re.search("\w+", x)]
        tokenized_text = [x for x in text if not x in stop_words]
    
        from nltk.stem import WordNetLemmatizer
        nltk.download('wordnet')
        lemmatizer = WordNetLemmatizer()
        tokenized_text = [lemmatizer.lemmatize(word) for word in tokenized_text]

        #Vectorizing
        from sklearn.feature_extraction.text import CountVectorizer
        count_vectorizer = CountVectorizer()
        vectorized_text = count_vectorizer.fit_transform(tokenized_text)
        feature_names = count_vectorizer.get_feature_names()
        vectorized_text_df = pd.DataFrame(vectorized_text.toarray(), columns = feature_names)
        vectorized_text_df.head()

        vectorized_text_df = [vectorized_text_df.sum()]
        vectorized_text_df = pd.DataFrame(vectorized_text_df)
    
        return(vectorized_text_df)
    
